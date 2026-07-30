"""Build the evidence-locked Vietnamese capstone report as a polished DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CAPSTONE_OS/04_OUTPUT/report/FINAL_REPORT_DRAFT.md"
OUTPUT = ROOT / "CAPSTONE_OS/04_OUTPUT/report/VIETNAMESE_ENTERPRISE_RAG_FINAL_REPORT.docx"

NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(151, 112, 20)
GRAY = RGBColor(90, 90, 90)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_box(paragraph, fill=LIGHT_FILL, border="D8DEE7"):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), border)
        borders.append(node)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def add_inline_markdown(paragraph, text, *, base_size=11, color=None, italic=False):
    text = text.replace("`", "")
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        run = paragraph.add_run(clean)
        set_run_font(run, size=base_size, color=color, bold=bold, italic=italic)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("VIETNAMESE ENTERPRISE RAG  |  EVIDENCE-LOCKED REPORT")
    set_run_font(run, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    add_page_number(paragraph)


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("BÁO CÁO ĐỒ ÁN"), size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("VIETNAMESE ENTERPRISE RAG"), size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    set_run_font(
        subtitle.add_run("Hybrid Retrieval, Fine-tuned Cross-Encoder"),
        size=14,
        color=DARK_BLUE,
        bold=True,
    )
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.space_after = Pt(28)
    set_run_font(
        subtitle2.add_run("và kiểm soát truy cập theo phòng ban"),
        size=14,
        color=DARK_BLUE,
        bold=True,
    )

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_after = Pt(72)
    set_run_font(
        status.add_run("FINAL REPORT DRAFT  |  EVIDENCE-LOCKED"),
        size=10,
        color=GRAY,
        bold=True,
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(8)
    set_run_font(
        note.add_run("Nội dung kỹ thuật đã khóa theo artifact kiểm chứng."),
        size=10.5,
        color=GRAY,
        italic=True,
    )
    note2 = doc.add_paragraph()
    note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        note2.add_run("Thông tin hành chính và định dạng bìa cần điền theo biểu mẫu chính thức của trường."),
        size=10,
        color=GRAY,
        italic=True,
    )
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Mục lục nội dung", level=1)
    sections = [
        "1. Giới thiệu",
        "2. Cơ sở lý thuyết",
        "3. Phương pháp",
        "4. Kiến trúc và triển khai",
        "5. Kết quả",
        "6. Thảo luận",
        "7. Giới hạn",
        "8. Hướng phát triển",
        "9. Kết luận",
        "Phụ lục A — Nguồn bằng chứng",
        "Phụ lục B — Claim boundary bắt buộc",
    ]
    for item in sections:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markdown(p, item)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    add_inline_markdown(
        note,
        "Ghi chú: số trang mục lục có thể cập nhật tự động trong Word sau khi áp dụng template nộp chính thức.",
        base_size=9.5,
        color=GRAY,
        italic=True,
    )
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        del rows[1]
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    if column_count == 2:
        widths = [int(TABLE_WIDTH_DXA * 0.61), TABLE_WIDTH_DXA - int(TABLE_WIDTH_DXA * 0.61)]
    elif column_count == 3:
        widths = [3900, 2200, 3260]
    elif column_count == 4:
        widths = [3400, 1980, 1980, 2000]
    else:
        base = TABLE_WIDTH_DXA // column_count
        widths = [base] * column_count
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)

    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            text = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            add_inline_markdown(
                paragraph,
                text,
                base_size=9.5,
                color=NAVY if row_index == 0 else None,
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                set_cell_fill(cell, HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    set_paragraph_box(paragraph)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    for line_index, line in enumerate(lines):
        if line_index:
            paragraph.add_run().add_break()
        set_run_font(paragraph.add_run(line), name="Consolas", size=8.5, color=NAVY)


def render_markdown(doc, markdown):
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines = []
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1:
                # Source title is represented by the cover.
                index += 1
                continue
            if level == 2 and (re.match(r"^\d+\.", text) or text.startswith("Phụ lục")):
                if doc.paragraphs and not doc.paragraphs[-1].text.endswith("Mục lục nội dung"):
                    doc.add_page_break()
                doc.add_heading(text, level=1)
            else:
                doc.add_heading(text, level=min(level - 1, 3))
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            set_paragraph_box(paragraph)
            paragraph.paragraph_format.left_indent = Inches(0.08)
            paragraph.paragraph_format.right_indent = Inches(0.08)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            add_inline_markdown(
                paragraph,
                stripped.lstrip("> ").strip(),
                base_size=10.5,
                color=DARK_BLUE,
                italic=True,
            )
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
            add_inline_markdown(paragraph, (numbered or bullet).group(1))
            index += 1
            continue

        if not stripped:
            flush_paragraph()
        else:
            paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    # The first three source lines are represented by the designed cover.
    marker = "## Tóm tắt"
    if marker not in markdown:
        raise RuntimeError("Expected report marker not found")
    body = markdown[markdown.index(marker) :]

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_page(section)
        configure_header_footer(section)

    add_cover(doc)
    add_toc(doc)
    render_markdown(doc, body)

    core = doc.core_properties
    core.title = "Vietnamese Enterprise RAG - Final Report Draft"
    core.subject = "Hybrid Retrieval, Fine-tuned Cross-Encoder và ACL/RLS"
    core.author = "Vietnamese Enterprise RAG Project Team"
    core.keywords = "Vietnamese RAG, BM25, Dense Retrieval, Cross-Encoder, RLS"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

"""
Document Loader Module
Hỗ trợ đọc các định dạng tài liệu khác nhau (PDF, Word, Text)
"""

import os
import csv
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

from src.utils import setup_logger, get_file_extension, is_supported_format
from config.constants import SUPPORTED_DOCUMENT_FORMATS

logger = setup_logger(__name__)


@dataclass
class Document:
    """Đại diện cho một tài liệu đã được tải"""
    
    content: str
    filename: str
    file_path: str
    file_type: str
    file_size: int
    loaded_at: str
    metadata: Dict[str, Any]
    
    def __repr__(self) -> str:
        return f"Document(filename='{self.filename}', size={self.file_size}, type='{self.file_type}')"


class DocumentLoader:
    """Loader cho các định dạng tài liệu khác nhau"""
    
    def __init__(self):
        """Khởi tạo DocumentLoader"""
        self.logger = logger
        self.supported_formats = SUPPORTED_DOCUMENT_FORMATS
    
    def load_document(self, file_path: str) -> Optional[Document]:
        """
        Tải một tài liệu từ đường dẫn file
        
        Args:
            file_path: Đường dẫn tới file
            
        Returns:
            Document object hoặc None nếu lỗi
        """
        file_path = Path(file_path)
        
        # Kiểm tra file tồn tại
        if not file_path.exists():
            self.logger.error(f"File không tồn tại: {file_path}")
            return None
        
        # Kiểm tra định dạng được hỗ trợ
        extension = get_file_extension(file_path.name)
        if not is_supported_format(file_path.name):
            self.logger.error(f"Định dạng không được hỗ trợ: {extension}")
            return None
        
        try:
            extra_metadata: Dict[str, Any] = {}
            # Đọc nội dung dựa trên định dạng
            if extension == ".pdf":
                content = self._load_pdf(file_path)
            elif extension in [".docx", ".doc"]:
                content = self._load_word(file_path)
            elif extension in [".txt", ".md"]:
                content = self._load_text(file_path)
            elif extension in [".csv", ".xlsx", ".xls"]:
                content, extra_metadata = self._load_spreadsheet(file_path)
            else:
                self.logger.error(f"Định dạng không được xử lý: {extension}")
                return None
            
            if content is None:
                return None
            
            # Tạo Document object
            file_size = file_path.stat().st_size
            document = Document(
                content=content,
                filename=file_path.name,
                file_path=str(file_path),
                file_type=self.supported_formats.get(extension, "Unknown"),
                file_size=file_size,
                loaded_at=datetime.now().isoformat(),
                metadata={
                    "extension": extension,
                    "file_size_bytes": file_size,
                    "character_count": len(content),
                    "word_count": len(content.split()),
                    **extra_metadata,
                }
            )
            
            self.logger.info(f"✅ Tài liệu tải thành công: {file_path.name}")
            return document
            
        except Exception as e:
            self.logger.error(f"❌ Lỗi khi tải tài liệu {file_path.name}: {str(e)}")
            return None
    
    def load_all_documents(self, directory: str) -> List[Document]:
        """
        Tải tất cả tài liệu từ một thư mục
        
        Args:
            directory: Đường dẫn thư mục
            
        Returns:
            Danh sách Document objects
        """
        directory = Path(directory)
        
        if not directory.exists():
            self.logger.error(f"Thư mục không tồn tại: {directory}")
            return []
        
        documents = []
        
        # Lặp qua tất cả file trong thư mục
        for file_path in directory.iterdir():
            if file_path.is_file():
                doc = self.load_document(str(file_path))
                if doc:
                    documents.append(doc)
        
        self.logger.info(f"✅ Tải xong {len(documents)} tài liệu từ {directory}")
        return documents
    
    def _load_pdf(self, file_path: Path) -> Optional[str]:
        """
        Tải nội dung từ file PDF
        
        Args:
            file_path: Đường dẫn file PDF
            
        Returns:
            Nội dung văn bản hoặc None nếu lỗi
        """
        try:
            import PyPDF2
            
            content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        content.append(text)
            
            return '\n'.join(content) if content else None
            
        except ImportError:
            self.logger.error("PyPDF2 chưa được cài đặt. Cài đặt: pip install PyPDF2")
            return None
        except Exception as e:
            self.logger.error(f"Lỗi khi đọc PDF: {str(e)}")
            return None
    
    def _load_word(self, file_path: Path) -> Optional[str]:
        """
        Tải nội dung từ file Word (.docx, .doc)
        
        Args:
            file_path: Đường dẫn file Word
            
        Returns:
            Nội dung văn bản hoặc None nếu lỗi
        """
        extension = get_file_extension(file_path.name)
        
        try:
            if extension == ".docx":
                return self._load_docx(file_path)
            elif extension == ".doc":
                return self._load_doc(file_path)
        except Exception as e:
            self.logger.error(f"Lỗi khi đọc file Word: {str(e)}")
            return None
    
    def _load_docx(self, file_path: Path) -> Optional[str]:
        """Tải file .docx"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            content = []
            
            # Đọc từ paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Đọc từ tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    if any(row_text):
                        content.append(' | '.join(row_text))
            
            return '\n'.join(content) if content else None
            
        except ImportError:
            self.logger.error("python-docx chưa được cài đặt. Cài đặt: pip install python-docx")
            return None
    
    def _load_doc(self, file_path: Path) -> Optional[str]:
        """Tải file .doc (legacy Word format)"""
        try:
            # Thử sử dụng python-docx (có thể hỗ trợ .doc)
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            content = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(content) if content else None
        except Exception:
            self.logger.warning(f"Không thể đọc .doc file: {file_path.name}. Thử dùng python-docx hoặc convert sang .docx")
            return None
    
    def _load_text(self, file_path: Path) -> Optional[str]:
        """
        Tải file text (.txt, .md)
        
        Args:
            file_path: Đường dẫn file text
            
        Returns:
            Nội dung văn bản hoặc None nếu lỗi
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content if content.strip() else None
        except UnicodeDecodeError:
            # Thử encoding khác
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return content if content.strip() else None
            except Exception as e:
                self.logger.error(f"Lỗi encoding khi đọc text file: {str(e)}")
                return None
        except Exception as e:
            self.logger.error(f"Lỗi khi đọc text file: {str(e)}")
            return None
    
    def _load_spreadsheet(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Load CSV/XLSX/XLS as structured text for RAG indexing."""
        extension = get_file_extension(file_path.name)
        if extension == ".csv":
            return self._load_csv(file_path)
        return self._load_excel(file_path)

    def _load_csv(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        rows: List[List[str]] = []
        for encoding in ["utf-8-sig", "utf-8", "latin-1"]:
            try:
                with open(file_path, "r", encoding=encoding, newline="") as f:
                    reader = csv.reader(f)
                    rows = [[cell.strip() for cell in row] for row in reader]
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.logger.error(f"Lá»—i khi Ä‘á»c CSV: {str(e)}")
                return None, {}

        if not rows:
            return None, {}

        return self._rows_to_structured_text(file_path.name, "CSV", rows), {
            "source_type": "spreadsheet",
            "sheet_names": ["CSV"],
            "row_count": max(0, len(rows) - 1),
            "table_count": 1,
        }

    def _load_excel(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        try:
            import pandas as pd
        except ImportError:
            self.logger.error("pandas/openpyxl chÆ°a Ä‘Æ°á»£c cÃ i Ä‘áº·t. CÃ i Ä‘áº·t: pip install pandas openpyxl")
            return None, {}

        try:
            sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
        except Exception as e:
            self.logger.error(f"Lá»—i khi Ä‘á»c Excel: {str(e)}")
            return None, {}

        sections = []
        total_rows = 0
        for sheet_name, df in sheets.items():
            df = df.fillna("")
            rows = [list(map(str, df.columns.tolist()))]
            rows.extend(df.astype(str).values.tolist())
            total_rows += len(df)
            sections.append(self._rows_to_structured_text(file_path.name, str(sheet_name), rows))

        content = "\n\n".join(section for section in sections if section.strip())
        if not content.strip():
            return None, {}

        return content, {
            "source_type": "spreadsheet",
            "sheet_names": [str(name) for name in sheets.keys()],
            "row_count": total_rows,
            "table_count": len(sheets),
        }

    def _rows_to_structured_text(self, filename: str, sheet_name: str, rows: List[List[str]]) -> str:
        """Convert tabular rows into compact row-oriented text with headers."""
        if not rows:
            return ""

        headers = [h.strip() or f"Column {idx + 1}" for idx, h in enumerate(rows[0])]
        lines = [f"File: {filename}", f"Sheet: {sheet_name}", f"Headers: {' | '.join(headers)}"]

        for row_index, row in enumerate(rows[1:], start=1):
            if not any(str(cell).strip() for cell in row):
                continue
            padded = list(row) + [""] * max(0, len(headers) - len(row))
            pairs = [
                f"{headers[col_index]}: {str(value).strip()}"
                for col_index, value in enumerate(padded[:len(headers)])
                if str(value).strip()
            ]
            if pairs:
                lines.append(f"Row {row_index}: " + "; ".join(pairs))

        return "\n".join(lines)

    def batch_load_documents(self, file_paths: List[str]) -> List[Document]:
        """
        Tải nhiều tài liệu từ danh sách đường dẫn
        
        Args:
            file_paths: Danh sách đường dẫn file
            
        Returns:
            Danh sách Document objects
        """
        documents = []
        for file_path in file_paths:
            doc = self.load_document(file_path)
            if doc:
                documents.append(doc)
        
        self.logger.info(f"✅ Tải xong {len(documents)}/{len(file_paths)} tài liệu")
        return documents


class DocumentManager:
    """Quản lý tập hợp các tài liệu"""
    
    def __init__(self):
        """Khởi tạo DocumentManager"""
        self.loader = DocumentLoader()
        self.documents: List[Document] = []
        self.logger = logger
    
    def add_document(self, file_path: str) -> bool:
        """
        Thêm một tài liệu
        
        Args:
            file_path: Đường dẫn file
            
        Returns:
            True nếu thành công, False nếu lỗi
        """
        doc = self.loader.load_document(file_path)
        if doc:
            self.documents.append(doc)
            return True
        return False
    
    def add_documents_from_directory(self, directory: str) -> int:
        """
        Thêm tất cả tài liệu từ một thư mục
        
        Args:
            directory: Đường dẫn thư mục
            
        Returns:
            Số lượng tài liệu được thêm
        """
        docs = self.loader.load_all_documents(directory)
        self.documents.extend(docs)
        return len(docs)
    
    def get_all_documents(self) -> List[Document]:
        """Lấy tất cả tài liệu"""
        return self.documents
    
    def get_document_by_filename(self, filename: str) -> Optional[Document]:
        """
        Tìm tài liệu theo tên file
        
        Args:
            filename: Tên file
            
        Returns:
            Document hoặc None
        """
        for doc in self.documents:
            if doc.filename == filename:
                return doc
        return None
    
    def clear_documents(self) -> None:
        """Xóa tất cả tài liệu"""
        self.documents.clear()
        self.logger.info("✅ Đã xóa tất cả tài liệu")
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        Lấy thống kê về các tài liệu
        
        Returns:
            Dictionary với thống kê
        """
        total_size = sum(doc.file_size for doc in self.documents)
        total_chars = sum(doc.metadata.get('character_count', 0) for doc in self.documents)
        total_words = sum(doc.metadata.get('word_count', 0) for doc in self.documents)
        
        file_types = {}
        for doc in self.documents:
            file_type = doc.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            "total_documents": len(self.documents),
            "total_size_bytes": total_size,
            "total_characters": total_chars,
            "total_words": total_words,
            "file_types": file_types,
        }

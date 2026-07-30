"""Small, deterministic file ingestion layer for the web product.

The original file is retained separately by the persistence backend.  This
module only extracts text and stable source locations for retrieval/citations.
It deliberately does not attempt OCR or legacy ``.doc`` conversion.
"""

from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import PurePath
from typing import Iterable


MAX_UPLOAD_BYTES = 15 * 1024 * 1024
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".txt",
    ".md",
    ".html",
    ".htm",
}


class IngestionError(ValueError):
    """An input file cannot be accepted or extracted safely."""


@dataclass(frozen=True)
class ExtractedChunk:
    content: str
    locator: str
    chunk_index: int


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    extension: str
    mime_type: str
    content: str
    chunks: list[ExtractedChunk]


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._ignored_depth = 0
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag in {"script", "style", "noscript"}:
            self._ignored_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._ignored_depth:
            self._ignored_depth -= 1
        if tag in {"p", "br", "li", "h1", "h2", "h3", "h4", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth:
            self.parts.append(data)


def extract_uploaded_document(
    *, filename: str, payload: bytes, mime_type: str = ""
) -> ExtractedDocument:
    """Extract supported office files into citation-ready chunks.

    The caller is expected to persist ``payload`` privately only after this
    validation succeeds.  No path supplied by a browser is trusted.
    """
    safe_name = PurePath(filename or "").name
    extension = PurePath(safe_name).suffix.lower()
    if not safe_name or extension not in SUPPORTED_EXTENSIONS:
        accepted = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise IngestionError(f"Định dạng chưa hỗ trợ. Chỉ nhận: {accepted}")
    if not payload:
        raise IngestionError("Tệp rỗng")
    if len(payload) > MAX_UPLOAD_BYTES:
        raise IngestionError("Tệp vượt giới hạn 15 MB")

    if extension == ".pdf":
        sections = _extract_pdf(payload)
    elif extension == ".docx":
        sections = _extract_docx(payload)
    elif extension in {".xlsx", ".xls"}:
        sections = _extract_excel(payload, extension)
    elif extension == ".csv":
        sections = _extract_csv(payload)
    elif extension in {".txt", ".md"}:
        sections = [("Nội dung", _decode_text(payload))]
    else:
        sections = [("Nội dung HTML", _extract_html(payload))]

    chunks = _sections_to_chunks(sections)
    if not chunks:
        raise IngestionError(
            "Không trích xuất được văn bản. PDF scan/ảnh cần OCR ở pha sau."
        )
    return ExtractedDocument(
        filename=safe_name,
        extension=extension,
        mime_type=mime_type or _default_mime(extension),
        content="\n\n".join(chunk.content for chunk in chunks),
        chunks=chunks,
    )


def _extract_pdf(payload: bytes) -> list[tuple[str, str]]:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency contract
        raise IngestionError("Máy chủ chưa có bộ đọc PDF") from exc
    reader = PdfReader(BytesIO(payload))
    return [
        (f"Trang {page_number}", page.extract_text() or "")
        for page_number, page in enumerate(reader.pages, start=1)
    ]


def _extract_docx(payload: bytes) -> list[tuple[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency contract
        raise IngestionError("Máy chủ chưa có bộ đọc DOCX") from exc
    document = Document(BytesIO(payload))
    sections: list[tuple[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text.strip():
            sections.append((f"Đoạn {index}", paragraph.text.strip()))
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                sections.append(
                    (f"Bảng {table_index}, hàng {row_index}", " | ".join(values))
                )
    return sections


def _extract_excel(payload: bytes, extension: str) -> list[tuple[str, str]]:
    try:
        import pandas as pd
    except ImportError as exc:  # pragma: no cover - dependency contract
        raise IngestionError("Máy chủ chưa có bộ đọc Excel") from exc
    engine = "xlrd" if extension == ".xls" else "openpyxl"
    try:
        sheets = pd.read_excel(BytesIO(payload), sheet_name=None, dtype=str, engine=engine)
    except Exception as exc:
        raise IngestionError("Không thể đọc workbook Excel") from exc
    sections: list[tuple[str, str]] = []
    for sheet_name, dataframe in sheets.items():
        frame = dataframe.fillna("")
        headers = [str(value).strip() or f"Cột {index + 1}" for index, value in enumerate(frame.columns)]
        for row_number, values in enumerate(frame.astype(str).values.tolist(), start=2):
            pairs = [
                f"{header}: {value.strip()}"
                for header, value in zip(headers, values, strict=False)
                if value.strip()
            ]
            if pairs:
                sections.append((f"Sheet {sheet_name}, hàng {row_number}", "; ".join(pairs)))
    return sections


def _extract_csv(payload: bytes) -> list[tuple[str, str]]:
    import csv

    text = _decode_text(payload)
    rows = list(csv.reader(text.splitlines()))
    if not rows:
        return []
    headers = [value.strip() or f"Cột {index + 1}" for index, value in enumerate(rows[0])]
    sections: list[tuple[str, str]] = []
    for row_number, values in enumerate(rows[1:], start=2):
        pairs = [
            f"{header}: {value.strip()}"
            for header, value in zip(headers, values, strict=False)
            if value.strip()
        ]
        if pairs:
            sections.append((f"CSV, hàng {row_number}", "; ".join(pairs)))
    return sections


def _extract_html(payload: bytes) -> str:
    parser = _VisibleTextParser()
    parser.feed(_decode_text(payload))
    return " ".join("".join(parser.parts).split())


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1258", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise IngestionError("Không nhận diện được bảng mã văn bản")


def _sections_to_chunks(sections: Iterable[tuple[str, str]]) -> list[ExtractedChunk]:
    chunks: list[ExtractedChunk] = []
    for locator, raw_text in sections:
        text = " ".join(raw_text.split())
        if not text:
            continue
        for piece in _split_text(text):
            chunks.append(ExtractedChunk(piece, locator, len(chunks)))
    return chunks


def _split_text(text: str, *, size: int = 1_200, overlap: int = 160) -> list[str]:
    if len(text) <= size:
        return [text]
    pieces: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        if end < len(text):
            boundary = text.rfind(" ", start + size // 2, end)
            if boundary > start:
                end = boundary
        pieces.append(text[start:end].strip())
        if end == len(text):
            break
        start = max(start + 1, end - overlap)
    return pieces


def _default_mime(extension: str) -> str:
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".html": "text/html",
        ".htm": "text/html",
    }[extension]
